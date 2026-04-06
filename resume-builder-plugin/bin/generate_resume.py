#!/usr/bin/env python3
"""
Generate an ATS-optimized 2-page resume in Word (.docx) format.

Usage: python3 generate_resume.py <input.json> <output.docx>

The input JSON schema:
{
  "name": "Full Name",
  "contact": {
    "phone": "123-456-7890",
    "email": "email@example.com",
    "linkedin": "linkedin.com/in/username"
  },
  "summary": "Professional summary paragraph...",
  "experience": [
    {
      "title": "Job Title",
      "company": "Company Name",
      "location": "City, ST",
      "dates": "Month Year - Month Year",
      "description": "Optional role context (1-2 sentences).",
      "bullets": ["Achievement 1", "Achievement 2"]
    }
  ],
  "additional_experience": [
    {
      "title": "Job Title",
      "company": "Company Name",
      "location": "City, ST",
      "dates": "Month Year - Month Year"
    }
  ],
  "education": [
    {
      "degree": "Degree Name",
      "school": "School Name"
    }
  ],
  "skills": ["Skill 1", "Skill 2"]
}
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# --- Design tokens ---
FONT = "Calibri"
COLOR_HEADING = RGBColor(0x1F, 0x2B, 0x4D)
COLOR_BODY = RGBColor(0x2D, 0x2D, 0x2D)
COLOR_SECONDARY = RGBColor(0x55, 0x55, 0x66)
COLOR_ACCENT = RGBColor(0x2C, 0x5F, 0x8A)
BORDER_HEX = "1F2B4D"


# --- Helpers ---

def _set_rfonts(run):
    """Ensure font family applies to all script types (ATS compatibility)."""
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rPr.insert(
        0,
        parse_xml(
            f'<w:rFonts {nsdecls("w")}'
            f' w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
        ),
    )


def fmt(run, size, bold=False, italic=False, color=None):
    """Apply font formatting to a run."""
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    _set_rfonts(run)
    return run


def sp(para, before=0, after=0, line=None):
    """Set paragraph spacing in points."""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)


def _bottom_border(para):
    """Add a thin bottom border to a paragraph (section separator)."""
    pPr = para._element.get_or_add_pPr()
    pPr.append(
        parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1"'
            f' w:color="{BORDER_HEX}"/>'
            f"</w:pBdr>"
        )
    )


# --- Section builders ---

def _heading(doc, text):
    """Add an uppercase section heading with bottom border."""
    p = doc.add_paragraph()
    fmt(p.add_run(text.upper()), 11, bold=True, color=COLOR_HEADING)
    sp(p, before=10, after=3)
    _bottom_border(p)
    return p


def _name_block(doc, data):
    """Centered name and contact line."""
    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p.add_run(data["name"]), 20, bold=True, color=COLOR_HEADING)
    sp(p, before=0, after=1)

    # Contact
    c = data["contact"]
    parts = [v for v in [c.get("phone"), c.get("email"), c.get("linkedin")] if v]
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(cp.add_run(" | ".join(parts)), 10, color=COLOR_SECONDARY)
    sp(cp, before=0, after=6)


def _summary(doc, text):
    """Professional summary section."""
    _heading(doc, "Professional Summary")
    p = doc.add_paragraph()
    fmt(p.add_run(text), 10, color=COLOR_BODY)
    sp(p, before=2, after=4)


def _experience_entry(doc, exp, condensed=False):
    """Single experience block: title, company line, optional description & bullets."""
    # Job title
    tp = doc.add_paragraph()
    fmt(tp.add_run(exp["title"]), 10.5, bold=True, color=COLOR_BODY)
    sp(tp, before=6, after=0)

    # Company | Location | Dates
    meta_parts = [exp["company"]]
    if exp.get("location"):
        meta_parts.append(exp["location"])
    meta_parts.append(exp["dates"])
    mp = doc.add_paragraph()
    fmt(mp.add_run(" | ".join(meta_parts)), 9.5, italic=True, color=COLOR_SECONDARY)
    sp(mp, before=0, after=2)

    if condensed:
        return

    # Role description
    if exp.get("description"):
        dp = doc.add_paragraph()
        fmt(dp.add_run(exp["description"]), 10, color=COLOR_BODY)
        sp(dp, before=1, after=2)

    # Bullets
    for bullet_text in exp.get("bullets", []):
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = Inches(0.2)
        bp.paragraph_format.first_line_indent = Inches(-0.15)
        fmt(bp.add_run("\u2022 "), 10, color=COLOR_ACCENT)
        fmt(bp.add_run(bullet_text), 10, color=COLOR_BODY)
        sp(bp, before=0.5, after=0.5, line=12.5)


def _experience_section(doc, data):
    """Professional Experience section with all roles."""
    _heading(doc, "Professional Experience")
    for exp in data.get("experience", []):
        _experience_entry(doc, exp)


def _additional_experience(doc, data):
    """Condensed older roles: title + company + dates only."""
    entries = data.get("additional_experience", [])
    if not entries:
        return
    _heading(doc, "Additional Experience")
    for exp in entries:
        _experience_entry(doc, exp, condensed=True)


def _education(doc, data):
    """Education section."""
    _heading(doc, "Education")
    for edu in data.get("education", []):
        p = doc.add_paragraph()
        fmt(p.add_run(edu["degree"]), 10, bold=True, color=COLOR_BODY)
        if edu.get("school"):
            fmt(p.add_run(" \u2014 "), 10, color=COLOR_SECONDARY)
            fmt(p.add_run(edu["school"]), 10, color=COLOR_SECONDARY)
        sp(p, before=2, after=2)


def _skills(doc, data):
    """Core Competencies section: dot-separated keyword list."""
    skills = data.get("skills", [])
    if not skills:
        return
    _heading(doc, "Core Competencies")
    p = doc.add_paragraph()
    fmt(p.add_run(" \u2022 ".join(skills)), 10, color=COLOR_BODY)
    sp(p, before=2, after=2)


# --- Main ---

def generate(data, output_path):
    """Build the complete resume document."""
    doc = Document()

    # Page setup — tight margins to maximize content in 2 pages
    sec = doc.sections[0]
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # Build sections
    _name_block(doc, data)
    _summary(doc, data["summary"])
    _experience_section(doc, data)
    _additional_experience(doc, data)
    _education(doc, data)
    _skills(doc, data)

    doc.save(output_path)
    print(f"Resume saved: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python3 generate_resume.py input.json output.docx")
        sys.exit(1)

    with open(sys.argv[1]) as f:
        data = json.load(f)
    generate(data, sys.argv[2])
