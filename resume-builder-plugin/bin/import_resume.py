#!/usr/bin/env python3
"""
Extract text from resume files for conversion to resume.md format.

Usage: python3 import_resume.py <resume_file>

Supported formats: .docx, .pdf, .txt, .md
Outputs extracted text to stdout for Claude to structure into resume.md.

Requires:
  - python-docx (for .docx): pip3 install python-docx
  - pdfplumber (for .pdf):   pip3 install pdfplumber
"""

import sys
from pathlib import Path


def extract_docx(path):
    """Extract text from .docx, handling table-based layouts and merged cells."""
    from docx import Document

    doc = Document(path)
    lines = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Tables (common in resume templates — merged cells repeat text)
    for table in doc.tables:
        for row in table.rows:
            seen_in_row = set()
            row_texts = []
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in seen_in_row:
                    row_texts.append(text)
                    seen_in_row.add(text)
            if row_texts:
                lines.append("\n".join(row_texts))

    return "\n".join(lines)


def extract_pdf(path):
    """Extract text from .pdf using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        print("Error: PDF support requires pdfplumber.", file=sys.stderr)
        print("Install with: pip3 install pdfplumber", file=sys.stderr)
        sys.exit(1)

    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n\n".join(parts)


def extract_text(path):
    """Read plain text or markdown files."""
    return Path(path).read_text(encoding="utf-8")


EXTRACTORS = {
    ".docx": extract_docx,
    ".pdf": extract_pdf,
    ".txt": extract_text,
    ".md": extract_text,
}


def main():
    if len(sys.argv) != 2:
        print("Usage: python3 import_resume.py <resume_file>")
        print(f"Supported formats: {', '.join(EXTRACTORS.keys())}")
        sys.exit(1)

    path = Path(sys.argv[1])

    if not path.exists():
        print(f"Error: File not found: {path}", file=sys.stderr)
        sys.exit(1)

    ext = path.suffix.lower()
    if ext not in EXTRACTORS:
        print(f"Unsupported format: {ext}", file=sys.stderr)
        print(f"Supported: {', '.join(EXTRACTORS.keys())}", file=sys.stderr)
        sys.exit(1)

    print(EXTRACTORS[ext](str(path)))


if __name__ == "__main__":
    main()
